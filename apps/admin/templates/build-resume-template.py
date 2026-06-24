"""
양식 docx 의 구조 + style 을 보존하면서 cell text 만 docxtemplater
placeholder 로 치환한 template 을 생성한다.

원본 sample (실제 고객 이력서) 의 PII 는 들어가지 않음.

핵심 원칙:
  - 각 paragraph 의 첫 run 의 text 만 placeholder 로 교체 → 폰트 크기/색/굵기 유지
  - 같은 paragraph 의 나머지 run text 는 비움 (분할된 run 들이 보존된 첫 run 의 끝에 빈 텍스트로 남음)
  - 학력/경력 같은 row-loop 표는 데이터 row 1 개만 남기고 docxtemplater table loop syntax 추가

사용법:
  python build-resume-template.py <원본 sample.docx>
  → 같은 폴더에 resume-template.docx 생성
"""

import sys
import os
from docx import Document
from docx.oxml.ns import qn


def clear_runs_keep_style(paragraph, new_text):
    """paragraph 의 첫 run 의 text 를 new_text 로 교체.
    같은 paragraph 의 나머지 run 들의 text 는 비움 (style 자체는 보존).
    """
    runs = paragraph.runs
    if not runs:
        # run 이 하나도 없으면 새로 생성 (style 없음)
        paragraph.add_run(new_text)
        return
    runs[0].text = new_text
    for r in runs[1:]:
        r.text = ""


def remove_extra_paragraphs(cell, keep_count):
    """cell 의 paragraph 중 첫 keep_count 개만 유지, 나머지 삭제.
    단 keep_count 보다 적게 있으면 추가하지 않음 (cell 의 inline drawing 보존 위해).
    """
    paragraphs = cell.paragraphs
    for p in paragraphs[keep_count:]:
        p._element.getparent().remove(p._element)


def remove_all_inline_drawings(cell):
    """cell 안의 모든 image (drawing) 요소를 제거. PII (실제 사진) 제거 용도."""
    body = cell._tc
    for drawing in body.iter(qn("w:drawing")):
        # parent run 까지 통째로 제거
        run_el = drawing.getparent()
        if run_el is not None and run_el.getparent() is not None:
            run_el.getparent().remove(run_el)


def set_cell_lines(cell, lines):
    """cell 안의 paragraph 별로 lines[i] 텍스트를 채움. paragraph 가 부족하면 새로 추가
    (style 없는 plain paragraph). 사용 후 lines 보다 많은 paragraph 는 잘림.
    """
    paragraphs = cell.paragraphs
    # 부족분 추가
    while len(paragraphs) < len(lines):
        cell.add_paragraph("")
        paragraphs = cell.paragraphs
    for i, text in enumerate(lines):
        clear_runs_keep_style(paragraphs[i], text)
    # 초과분 제거
    for p in paragraphs[len(lines):]:
        p._element.getparent().remove(p._element)


def reset_row_to_single_data(table, header_rows=1):
    """table 의 데이터 row 들 (header 이후) 을 1개만 남기고 모두 제거."""
    for r in list(table.rows[header_rows + 1:]):
        table._element.remove(r._element)


def main():
    if len(sys.argv) < 2:
        print("usage: python build-resume-template.py <sample.docx>")
        sys.exit(1)
    src = sys.argv[1]
    dst = os.path.join(os.path.dirname(__file__), "resume-template.docx")

    doc = Document(src)
    tables = doc.tables

    # ===== 헤더 (table 0) =====
    # [0,0] = 텍스트 (이름·생년월일·전화·이메일·주소·한 줄 자기소개) — paragraph 별 폰트 크기 다름
    # [0,1] = 사진 (image inline)
    h = tables[0]
    # [0,0]: 4 줄 (이름 / 정보 / 주소 / 한 줄) — 원본 paragraph 의 style 유지
    set_cell_lines(
        h.rows[0].cells[0],
        [
            "{name_vi} / {name_kr}",
            "{birth_date}  |  {phone}  |  {email}",
            "{address}",
            "{one_liner}",
        ],
    )

    # [0,1]: 사진 PII 제거 + image placeholder
    remove_all_inline_drawings(h.rows[0].cells[1])
    # cell 의 paragraph 들 모두 비우고 첫 paragraph 에 {%photo} 넣기
    set_cell_lines(h.rows[0].cells[1], ["{%photo}"])

    # ===== 학력 (table 2) =====
    edu = tables[2]
    reset_row_to_single_data(edu, header_rows=1)
    row = edu.rows[1]
    clear_runs_keep_style(row.cells[0].paragraphs[0], "{#educations}{school}")
    clear_runs_keep_style(row.cells[1].paragraphs[0], "{major}")
    clear_runs_keep_style(row.cells[2].paragraphs[0], "{period}")
    clear_runs_keep_style(row.cells[3].paragraphs[0], "{status}{/educations}")
    for c in row.cells:
        remove_extra_paragraphs(c, 1)

    # ===== 경력 (table 4) =====
    car = tables[4]
    reset_row_to_single_data(car, header_rows=1)
    row = car.rows[1]
    clear_runs_keep_style(row.cells[0].paragraphs[0], "{#careers}{workplace}")
    clear_runs_keep_style(row.cells[1].paragraphs[0], "{period}")
    clear_runs_keep_style(row.cells[2].paragraphs[0], "{role}")
    clear_runs_keep_style(row.cells[3].paragraphs[0], "{detail}")
    clear_runs_keep_style(row.cells[4].paragraphs[0], "{status}{/careers}")
    for c in row.cells:
        remove_extra_paragraphs(c, 1)

    # ===== 자격증·수상 (table 6) =====
    cert = tables[6]
    reset_row_to_single_data(cert, header_rows=1)
    row = cert.rows[1]
    clear_runs_keep_style(row.cells[0].paragraphs[0], "{#certifications}{name}")
    clear_runs_keep_style(row.cells[1].paragraphs[0], "{issuer}")
    clear_runs_keep_style(row.cells[2].paragraphs[0], "{date}{/certifications}")
    for c in row.cells:
        remove_extra_paragraphs(c, 1)

    # ===== 기술·어학 (table 8) =====
    skill = tables[8]
    reset_row_to_single_data(skill, header_rows=1)
    row = skill.rows[1]
    clear_runs_keep_style(row.cells[0].paragraphs[0], "{#skills}{name}")
    clear_runs_keep_style(row.cells[1].paragraphs[0], "{detail}")
    clear_runs_keep_style(row.cells[2].paragraphs[0], "{level}{/skills}")
    for c in row.cells:
        remove_extra_paragraphs(c, 1)

    # ===== 기타 활동 (table 10) =====
    other = tables[10]
    reset_row_to_single_data(other, header_rows=1)
    row = other.rows[1]
    clear_runs_keep_style(row.cells[0].paragraphs[0], "{#activities}{name}")
    clear_runs_keep_style(row.cells[1].paragraphs[0], "{period}")
    clear_runs_keep_style(row.cells[2].paragraphs[0], "{org}")
    clear_runs_keep_style(row.cells[3].paragraphs[0], "{detail}{/activities}")
    for c in row.cells:
        remove_extra_paragraphs(c, 1)

    # ===== 자기소개 본문 (table 12) =====
    intro = tables[12]
    intro_cell = intro.rows[0].cells[0]
    # 첫 paragraph 의 style 만 보존, 나머지 paragraph 제거, 첫 run text → {narrative}
    # docxtemplater 의 linebreaks:true 옵션이 \n 을 줄바꿈으로 변환.
    if intro_cell.paragraphs:
        clear_runs_keep_style(intro_cell.paragraphs[0], "{narrative}")
    remove_extra_paragraphs(intro_cell, 1)

    doc.save(dst)
    print(f"saved: {dst}")


if __name__ == "__main__":
    main()
